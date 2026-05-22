import io

from docx import Document


def extract_docx(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text)
            if row_text:
                chunks.append(row_text)

    return "\n".join(chunks)
